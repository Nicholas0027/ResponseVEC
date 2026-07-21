#!/usr/bin/env python
"""Generate a standalone Colab notebook for the two extra comparators requested
after the main pipeline already ran: BGE (sentence encoder) and RespondentVec
(query-independent ablation, design §2.3.D).

Assumes the MAIN pipeline notebook has already been run at least through Cell 8
(train_primary) on Drive:/ResponseVec_colab — this notebook reuses the existing
runtime_config.yaml, item_folds.json, population prior, and option_table.npz
from that run rather than rebuilding them. It follows the same
setup/patch/report conventions as the main notebook so the same Drive folder
and Claude<->Drive loop apply unchanged.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path


def markdown(source: str) -> dict:
    return {"cell_type": "markdown", "metadata": {}, "source": source.splitlines(keepends=True)}


def code(source: str) -> dict:
    return {"cell_type": "code", "execution_count": None, "metadata": {}, "outputs": [], "source": source.splitlines(keepends=True)}


INTRO = """# ResponseVec extras — BGE (sentence) + RespondentVec, standalone

Run this AFTER the main pipeline notebook has completed through Cell 8
(`train_primary.py`) at least once on Drive:/ResponseVec_colab — this notebook
reuses that run's `runtime_config.yaml`, `item_folds.json`, population prior,
and `option_table.npz` rather than rebuilding them.

**Two additions, both cheap compared to the 8B extraction already done:**
1. **BGE (`sentence` family)** — a small, generic off-the-shelf text embedder
   (BAAI/bge-base-en-v1.5), mirroring LLMGeovec's Bert-whitening/GTE controls:
   does the TASK model's own hidden states beat a generic embedder, or is any
   embedder just as good? Already wired into `train_primary.py`; just needs
   extraction.
2. **RespondentVec** (design §2.3.D) — a query-independent respondent vector
   r_i^K (target question and options removed entirely), encoded with the same
   input-centric LLM2Vec checkpoint. Backs H9 and the "a vector cannot
   represent a person" rebuttal. New code this session: `build_respondent_prompt`,
   `build_respondentvec_units`, `arrays_from_respondent_cache`,
   `scripts/extract_respondentvec.py`, and a `--include-respondent-vec` flag on
   `train_primary.py`.

Both are added as EXTRA comparators — they do not change the primary G1 gate
(response_centric vs direct/input_centric/raw), which was already decided by
the main pipeline run.
"""

SETUP = r'''# === Cell 1: setup (same Drive folder + patch loop as the main notebook) ===
import os, shutil, subprocess, zipfile, glob, sys, json
from pathlib import Path

IN_COLAB = "google.colab" in sys.modules
if IN_COLAB:
    from google.colab import drive
    drive.mount('/content/drive', force_remount=False)
    DRIVE = '/content/drive/MyDrive/ResponseVec_colab'
    RESULTS = f'{DRIVE}/results'
    BACKUP = f'{DRIVE}/artifacts_backup'
    os.makedirs(RESULTS, exist_ok=True); os.makedirs(BACKUP, exist_ok=True)

    zips = sorted(glob.glob(f'{DRIVE}/responsevec*.zip'), key=os.path.getmtime)
    assert zips, f'No responsevec.zip found in {DRIVE} — drag the latest one in first.'
    print('using code package:', zips[-1])
    shutil.rmtree('/content/responsevec', ignore_errors=True)
    with zipfile.ZipFile(zips[-1]) as z: z.extractall('/content')
    PROJECT_ROOT = Path('/content/responsevec')
    ARTIFACT_ROOT = Path(BACKUP) / 'artifacts'

    os.makedirs(f'{DRIVE}/patches', exist_ok=True)
    for patch in sorted(glob.glob(f'{DRIVE}/patches/patch_*.py')):
        print('[patch]', os.path.basename(patch))
        exec(compile(open(patch).read(), patch, 'exec'), {'RESPONSEVEC_ROOT': '/content/responsevec'})
else:
    PROJECT_ROOT = Path.cwd()
    if PROJECT_ROOT.name != 'responsevec':
        found = list(PROJECT_ROOT.rglob('responsevec/pyproject.toml'))
        assert found, 'Run from the responsevec project or upload the bundle in Colab'
        PROJECT_ROOT = found[0].parent
    ARTIFACT_ROOT = PROJECT_ROOT / 'artifacts' / 'notebook'
    DRIVE = RESULTS = BACKUP = None

os.chdir(PROJECT_ROOT)
subprocess.check_call([sys.executable, '-m', 'pip', 'install', '-q', '-e', '.[llm,llm2vec,analysis]'])
if IN_COLAB:
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', '-q', 'python-docx'])
    try:
        from google.colab import userdata
        os.environ['HF_TOKEN'] = userdata.get('HF_TOKEN')
        print('HF_TOKEN set')
    except Exception:
        print('no HF_TOKEN: gated models will fall back per config')

import torch
print('GPU:', torch.cuda.get_device_name(0) if torch.cuda.is_available() else 'NONE — connect a GPU runtime!')

RUNTIME_CONFIG = ARTIFACT_ROOT / 'runtime_config.yaml'
assert RUNTIME_CONFIG.exists(), f'{RUNTIME_CONFIG} not found — run the MAIN pipeline notebook through Cell 2 first.'
print('using existing runtime config:', RUNTIME_CONFIG)

def run(*parts):
    command = [str(p) for p in parts]
    print('+', ' '.join(command), flush=True)
    subprocess.check_call(command, cwd=PROJECT_ROOT)

def run_stage(name, parts):
    command = [str(p) for p in parts]
    print(f'\\n=== {name}: {" ".join(command)} ===', flush=True)
    log_path = PROJECT_ROOT / f'{name}.log'
    with open(log_path, 'w') as log:
        proc = subprocess.Popen(command, cwd=PROJECT_ROOT, stdout=subprocess.PIPE,
                                stderr=subprocess.STDOUT, text=True)
        for line in proc.stdout:
            print(line, end=''); log.write(line)
        proc.wait()
    sync(name, ok=(proc.returncode == 0))
    if proc.returncode != 0:
        raise RuntimeError(f'{name} FAILED (exit {proc.returncode}) — log synced to Drive for Claude')
    print(f'=== {name}: OK ===')

def sync(stage, ok=True):
    if not IN_COLAB:
        return
    import glob as _glob, shutil as _shutil
    for log in _glob.glob(str(PROJECT_ROOT / '*.log')):
        _shutil.copy(log, RESULTS)
    for csv in _glob.glob(str(ARTIFACT_ROOT / '**' / '*.csv'), recursive=True):
        os.makedirs(f'{RESULTS}/raw', exist_ok=True); _shutil.copy(csv, f'{RESULTS}/raw')
    make_report(stage, ok)

def make_report(stage, ok):
    from docx import Document
    import pandas as pd, datetime, glob as _glob
    doc = Document()
    doc.add_heading('ResponseVec extras report (BGE + RespondentVec)', 0)
    doc.add_paragraph(f'updated: {datetime.datetime.now().isoformat()}  |  last stage: {stage}  |  status: {"OK" if ok else "FAILED"}')
    root = str(ARTIFACT_ROOT)
    for csv in sorted(_glob.glob(f'{root}/**/*.csv', recursive=True)):
        doc.add_heading(os.path.basename(csv), level=2)
        try:
            frame = pd.read_csv(csv)
            doc.add_paragraph(f'{len(frame)} rows')
            doc.add_paragraph(frame.head(60).to_string())
        except Exception as exc:
            doc.add_paragraph(f'unreadable: {exc}')
    for log in sorted(_glob.glob(str(PROJECT_ROOT / '*.log'))):
        doc.add_heading(os.path.basename(log), level=2)
        doc.add_paragraph(''.join(open(log).readlines()[-60:]))
    doc.save(f'{RESULTS}/ResponseVec_report.docx')
    print(f'[report] {RESULTS}/ResponseVec_report.docx updated')

print('setup complete')
'''

CONFIG = r'''# === Cell 2: run parameters (match the main pipeline's primary run) ===
import yaml
K = 5
OUTER_FOLDS = list(range(6))
DECODER_SEEDS = [1701, 7, 42]
base = yaml.safe_load(RUNTIME_CONFIG.read_text())
print({"backbone": base["representation"]["backbone"],
       "sentence_encoder": base["representation"]["sentence_encoder"],
       "input_centric": base["representation"]["input_centric"]})
'''

EXTRACT_SENTENCE = r'''# === Cell 3: extract BGE (sentence) representations, K=5, all three splits ===
# Small model (BAAI/bge-base-en-v1.5) — much faster than the 8B/LLM2Vec passes
# already done. Test split gets the same 3 option-permutation seeds as every
# other family so it can be averaged the same way at evaluation time.
for split, option_seeds in (("train", [0]), ("validation", [0]), ("test", [0, 7, 42])):
    for option_seed in option_seeds:
        run_stage(f'extract_sentence_{split}_o{option_seed}',
            [sys.executable, 'scripts/extract_representations.py', '--config', RUNTIME_CONFIG,
             '--family', 'sentence', '--split', split, '--k', str(K),
             '--option-seed', str(option_seed), '--selection', 'semantic'])
'''

EXTRACT_RESPONDENTVEC = r'''# === Cell 4: extract RespondentVec (query-independent, design section 2.3.D) ===
# One vector per (panel_id, domain) -- not per target item -- so this is much
# cheaper than any query-conditioned family: no option_seed axis, and roughly
# (num_target_items_per_respondent)x fewer prompts to encode.
for split in ("train", "validation", "test"):
    run_stage(f'extract_respondentvec_{split}',
        [sys.executable, 'scripts/extract_respondentvec.py', '--config', RUNTIME_CONFIG,
         '--split', split, '--k', str(K)])
'''

TRAIN = r'''# === Cell 5: retrain all six outer-fold heads WITH the two new families ===
# Re-running train_primary.py is required (not just for the two new methods):
# it is the one script that reads every cached family and writes the combined
# predictions_seed_averaged.parquet per fold. --include-respondent-vec adds
# RespondentVec on top of the families already in representation_families
# (which already includes "sentence" from this session's earlier change).
option_seed_arg = ','.join(map(str, [0, 7, 42]))
decoder_seed_arg = ','.join(map(str, DECODER_SEEDS))
for fold in OUTER_FOLDS:
    run_stage(f'train_fold{fold}_k{K}_extras',
        [sys.executable, 'scripts/train_primary.py', '--config', RUNTIME_CONFIG,
         '--fold', str(fold), '--k', str(K),
         '--option-seeds', option_seed_arg, '--seeds', decoder_seed_arg,
         '--include-respondent-vec'])
'''

EVALUATE = r'''# === Cell 6: re-evaluate (adds the secondary raw-vs-sentence check; response_centric stays primary) ===
run_stage('evaluate_primary_extras', [sys.executable, 'scripts/evaluate_primary.py', '--config', RUNTIME_CONFIG, '--k', str(K)])
import pandas as pd, json as _json
metric_dir = ARTIFACT_ROOT / 'metrics' / f'k_{K}'
gate = _json.loads((metric_dir / 'gate_g1.json').read_text())
primary = pd.read_csv(metric_dir / 'primary_family_nll.csv')
display(primary); display(gate)
secondary_path = metric_dir / 'secondary_raw_vs_sentence.csv'
if secondary_path.exists():
    display(pd.read_csv(secondary_path))
table1 = pd.read_csv(metric_dir / 'table1_unseen.csv') if (metric_dir / 'table1_unseen.csv').exists() else None
if table1 is not None:
    display(table1[table1['method'].isin(['response_centric', 'raw_selected', 'input_centric', 'sentence', 'respondent_vec'])])
'''


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--notebook", default="output/jupyter-notebook/responsevec_extras.ipynb")
    args = parser.parse_args()
    project = Path(__file__).resolve().parents[1]
    path = (project / args.notebook).resolve()
    path.parent.mkdir(parents=True, exist_ok=True)
    notebook = {"cells": [], "metadata": {}, "nbformat": 4, "nbformat_minor": 5}

    notebook["cells"] = [
        markdown(INTRO),
        markdown("## 1. Setup — same Drive folder + patch loop as the main notebook"),
        code(SETUP),
        markdown("## 2. Run parameters"),
        code(CONFIG),
        markdown("## 3. Extract BGE (sentence) representations"),
        code(EXTRACT_SENTENCE),
        markdown("## 4. Extract RespondentVec (query-independent)"),
        code(EXTRACT_RESPONDENTVEC),
        markdown("## 5. Retrain all six outer-fold heads with both new families"),
        code(TRAIN),
        markdown("## 6. Re-evaluate — primary claim unchanged, secondary check + full table"),
        code(EVALUATE),
    ]
    for index, cell in enumerate(notebook["cells"]):
        cell["id"] = f"extras-{index:02d}"
    notebook["metadata"]["kernelspec"] = {"display_name": "Python 3", "language": "python", "name": "python3"}
    path.write_text(json.dumps(notebook, indent=1), encoding="utf-8")
    print(path)


if __name__ == "__main__":
    main()
