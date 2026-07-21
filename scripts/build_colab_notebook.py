#!/usr/bin/env python
"""Generate the one-click ResponseVec Colab pipeline notebook.

The notebook's contract with the Claude<->Drive collaboration loop (ported from
the RAPL pipeline so the human only ever runs cells):
- Cell 1 always re-extracts the freshest responsevec*.zip from Drive:/ResponseVec_colab,
  so a code fix pushed as a new zip takes effect on the next "Run cell".
- Claude-pushed patches in Drive:/ResponseVec_colab/patches/patch_*.py are applied
  in order every Cell-1 run, so most code fixes need no zip re-upload at all.
- Every stage syncs artifacts to Drive:/ResponseVec_colab/artifacts_backup (resume
  across runtime disconnects) and regenerates Drive:/ResponseVec_colab/results/
  ResponseVec_report.docx — a docx because that is what Claude's Drive reader can
  ingest; raw CSVs are copied alongside for the human.
- The scientific contract is unchanged: RUN_MODE="smoke" first; only "a100" spends
  GPU hours, and the E0/G0 signal audit (Cell 5) gates that spend.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path


def markdown(source: str) -> dict:
    return {"cell_type": "markdown", "metadata": {}, "source": source.splitlines(keepends=True)}


def code(source: str) -> dict:
    return {"cell_type": "code", "execution_count": None, "metadata": {}, "outputs": [], "source": source.splitlines(keepends=True)}


INTRO = """# ResponseVec AAAI experiment pipeline — one-click Colab runner

**Question.** Does a response-centric LLM representation predict a held-out real respondent answer better than the same model's direct output, a raw hidden-state probe, and input-centric LLM2Vec on entirely unseen survey items?

**Primary success rule.** At unseen-item K=5, ResponseVec must beat all three preregistered controls after Holm correction, with at least a 0.02-nat NLL improvement. This notebook never treats smoke outputs as evidence.

**准备（只需一次）**
1. Runtime → Change runtime type → **GPU**（A100 40GB+ 才能跑 bf16 8B 主实验；smoke 模式纯 CPU 即可）。
2. 左侧 🔑 Secrets 里添加（可选）：`HF_TOKEN`（跑 gated 模型；不填自动退回 ungated 骨干）。
3. 确认 Drive 的 `ResponseVec_colab/` 里有最新 `responsevec.zip`。

**运行顺序**：Cell 1（setup，永远安全重跑，自动取最新 zip + 打补丁 + 从 Drive 续跑）→ 2（runtime 检查）→ 3（CPU smoke，必过）→ 4（准备真实 SocioBench）→ **5（E0/G0 信号审计，CPU 出闸门判定 —— 先和 Claude 一起读 verdicts 再花 GPU）** → 6（选项编码，一次性）→ 7（K=5 表示抽取，最长）→ 8（六折头训练）→ 9（G1 主判定）→ 10（K 曲线，可选）→ 11（归档）。断线后重开 runtime，从 Cell 1 重跑即可 —— 所有阶段都从 Drive 备份自动续跑。每个阶段结束会更新 `ResponseVec_colab/results/ResponseVec_report.docx`，Claude 从那里读结果。跑到报错或出结果，告诉 Claude「读报告」即可。
"""

MODE = '''# === Cell 0: mode and execution contract ===
# Leave RUN_MODE="smoke" for the first execution. Change to "a100" only after
# smoke passes AND an A100 runtime is connected. The primary K=5 experiment runs
# before K curves or optional extensions.
RUN_MODE = "smoke"                  # "smoke" | "a100"
PRIMARY_K = 5
OPTION_SEEDS = [0, 7, 42]           # identity + two deterministic semantic permutations
DECODER_SEEDS = [1701, 7, 42]
OUTER_FOLDS = list(range(6))
RUN_K_CURVE = False                 # turn on only after G1 passes
K_VALUES = [0, 1, 3, 5, 8]

import sys
IN_COLAB = "google.colab" in sys.modules
print({"mode": RUN_MODE, "in_colab": IN_COLAB})
'''

SETUP = r'''# === Cell 1: setup (always safe to re-run; picks up the newest responsevec.zip) ===
import os, shutil, subprocess, zipfile, glob, sys, json
from pathlib import Path

if IN_COLAB:
    from google.colab import drive
    drive.mount('/content/drive', force_remount=False)
    DRIVE = '/content/drive/MyDrive/ResponseVec_colab'
    RESULTS = f'{DRIVE}/results'
    BACKUP = f'{DRIVE}/artifacts_backup'
    os.makedirs(RESULTS, exist_ok=True); os.makedirs(BACKUP, exist_ok=True)

    zips = sorted(glob.glob(f'{DRIVE}/responsevec*.zip'), key=os.path.getmtime)
    assert zips, f'No responsevec.zip found in {DRIVE} — drag it into the Drive folder first.'
    print('using code package:', zips[-1])
    # The bundle nests every file under a top-level responsevec/ directory, so
    # extracting to /content yields /content/responsevec/... — extract there,
    # clearing any prior copy first so a new zip fully replaces the old code.
    shutil.rmtree('/content/responsevec', ignore_errors=True)
    with zipfile.ZipFile(zips[-1]) as z: z.extractall('/content')
    PROJECT_ROOT = Path('/content/responsevec')
    ARTIFACT_ROOT = Path(BACKUP) / 'artifacts'

    # Apply Claude-pushed patches in order: whenever Claude fixes a bug it uploads
    # a patch to ResponseVec_colab/patches/; re-running this cell picks every patch
    # up, so no zip re-upload is ever needed for code fixes. Patches are idempotent.
    os.makedirs(f'{DRIVE}/patches', exist_ok=True)
    for patch in sorted(glob.glob(f'{DRIVE}/patches/patch_*.py')):
        print('[patch]', os.path.basename(patch))
        exec(compile(open(patch).read(), patch, 'exec'), {'RESPONSEVEC_ROOT': '/content/responsevec'})

    # Restore prior artifacts from Drive so every stage resumes instead of restarting.
    if os.path.exists(f'{BACKUP}/artifacts'):
        print('resuming from Drive backup:', f'{BACKUP}/artifacts')
else:
    PROJECT_ROOT = Path.cwd()
    if PROJECT_ROOT.name != 'responsevec':
        found = list(PROJECT_ROOT.rglob('responsevec/pyproject.toml'))
        assert found, 'Run from the responsevec project or upload the bundle in Colab'
        PROJECT_ROOT = found[0].parent
    ARTIFACT_ROOT = PROJECT_ROOT / 'artifacts' / 'notebook'
    DRIVE = RESULTS = BACKUP = None

os.chdir(PROJECT_ROOT)
ARTIFACT_ROOT.mkdir(parents=True, exist_ok=True)

extras = '.[dev]' if RUN_MODE == 'smoke' else '.[llm,llm2vec,analysis]'
subprocess.check_call([sys.executable, '-m', 'pip', 'install', '-q', '-e', extras])
if IN_COLAB:
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', '-q', 'python-docx'])
if RUN_MODE == 'a100':
    # The inference package's published deps pull training-only MTEB/W&B/Hydra,
    # FlashAttention, and an exact Torch build; the curated runtime already
    # supplies every import used for inference.
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', '-q', '--no-deps', 'llm2vec-gen==0.1.3'])

# Gated-model fallback: no HF_TOKEN -> leave config backbone (already ungated-safe).
if IN_COLAB:
    try:
        from google.colab import userdata
        os.environ['HF_TOKEN'] = userdata.get('HF_TOKEN')
        print('HF_TOKEN set')
    except Exception:
        print('no HF_TOKEN: gated models will fall back per config')

try:
    import torch
    print('GPU:', torch.cuda.get_device_name(0) if torch.cuda.is_available() else 'NONE — switch runtime to GPU for a100 mode!')
except Exception as exc:
    print('torch not importable yet:', exc)
print({"project": str(PROJECT_ROOT), "artifacts": str(ARTIFACT_ROOT)})
'''
HELPERS = r'''# === Cell 1b: stage runner + Drive sync + docx report (the Claude<->Drive loop) ===
import subprocess, glob, os, shutil, json
from pathlib import Path

def run(*parts):
    """Run a pipeline command inline (raises on failure); for quick checks."""
    command = [str(p) for p in parts]
    print('+', ' '.join(command), flush=True)
    subprocess.check_call(command, cwd=PROJECT_ROOT)

def run_stage(name, parts):
    """Run a pipeline command, tee output to a log, sync everything to Drive.
    On failure the log is synced first so Claude can read the traceback."""
    command = [str(p) for p in parts]
    print(f'\n=== {name}: {" ".join(command)} ===', flush=True)
    log_path = PROJECT_ROOT / f'{name}.log'
    with open(log_path, 'w') as log:
        proc = subprocess.Popen(command, cwd=PROJECT_ROOT, stdout=subprocess.PIPE,
                                stderr=subprocess.STDOUT, text=True)
        for line in proc.stdout:
            print(line, end=''); log.write(line)
        proc.wait()
    sync(name, ok=(proc.returncode == 0))
    if proc.returncode != 0:
        raise RuntimeError(f'{name} FAILED (exit {proc.returncode}) — log synced to Drive for Claude')
    print(f'=== {name}: OK ===')

def sync(stage, ok=True):
    if not IN_COLAB:
        return
    # ARTIFACT_ROOT already lives on Drive (artifacts_backup); just copy logs +
    # CSVs to results/ and rebuild the report Claude reads.
    for log in glob.glob(str(PROJECT_ROOT / '*.log')):
        shutil.copy(log, RESULTS)
    for csv in glob.glob(str(ARTIFACT_ROOT / '**' / '*.csv'), recursive=True):
        os.makedirs(f'{RESULTS}/raw', exist_ok=True); shutil.copy(csv, f'{RESULTS}/raw')
    make_report(stage, ok)

def make_report(stage, ok):
    """One docx with every audit/gate JSON + metric table + log tails — the file
    Claude reads from Drive (its reader ingests docx, not raw CSV)."""
    from docx import Document
    import pandas as pd, datetime, json as _json
    doc = Document()
    doc.add_heading('ResponseVec experiment report', 0)
    doc.add_paragraph(f'updated: {datetime.datetime.now().isoformat()}  |  last stage: {stage}  '
                      f'|  status: {"OK" if ok else "FAILED"}  |  mode: {RUN_MODE}')
    root = str(ARTIFACT_ROOT)
    # Decision-critical JSON first: signal audit (G0) and primary gate (G1).
    for pattern in ('**/signal_audit.json', '**/gate_g1.json', '**/*audit*.json', '**/smoke_summary.json'):
        for path in sorted(glob.glob(f'{root}/{pattern}', recursive=True)):
            doc.add_heading(path.replace(root, 'artifacts'), level=2)
            doc.add_paragraph(open(path).read()[:3000])
    # Primary-family tables and any other metric CSVs.
    for csv in sorted(glob.glob(f'{root}/**/*.csv', recursive=True)):
        doc.add_heading(os.path.basename(csv), level=2)
        try:
            frame = pd.read_csv(csv)
            doc.add_paragraph(f'{len(frame)} rows')
            doc.add_paragraph(frame.head(60).to_string())
        except Exception as exc:
            doc.add_paragraph(f'unreadable: {exc}')
    # Head-fit selection metadata (which direct/raw variant validation picked).
    for sel in sorted(glob.glob(f'{root}/**/direct_selection.json', recursive=True)) + \
               sorted(glob.glob(f'{root}/**/raw_selection.json', recursive=True)):
        doc.add_heading(sel.replace(root, 'artifacts'), level=2)
        doc.add_paragraph(open(sel).read()[:1500])
    for log in sorted(glob.glob(str(PROJECT_ROOT / '*.log'))):
        doc.add_heading(os.path.basename(log), level=2)
        doc.add_paragraph(''.join(open(log).readlines()[-60:]))
    doc.save(f'{RESULTS}/ResponseVec_report.docx')
    print(f'[report] {RESULTS}/ResponseVec_report.docx updated')

print('stage runner ready')
'''
RUNTIME = r'''# === Cell 2: runtime config (redirect all artifact paths onto Drive) ===
import yaml, json
base = yaml.safe_load((PROJECT_ROOT / 'configs/responsevec.yaml').read_text())
for key in ('processed', 'cache', 'heads', 'predictions', 'metrics', 'figures'):
    if key in base['paths']:
        base['paths'][key] = str(ARTIFACT_ROOT / key)
base['paths']['sociobench_repo'] = str(PROJECT_ROOT / 'external' / 'SocioBench')
RUNTIME_CONFIG = ARTIFACT_ROOT / 'runtime_config.yaml'
RUNTIME_CONFIG.write_text(yaml.safe_dump(base, sort_keys=False))
print('runtime config ->', RUNTIME_CONFIG)

if RUN_MODE == 'a100':
    import torch
    assert torch.cuda.is_available(), 'Connect an A100 GPU runtime before continuing'
    props = torch.cuda.get_device_properties(0)
    vram_gb = props.total_memory / 2**30
    assert vram_gb >= 30, f'Primary bf16 run requires >=30 GB VRAM; found {vram_gb:.1f} GB'
    print({"gpu": props.name, "vram_gb": round(vram_gb, 1)})
    run_stage('model_preflight', [sys.executable, 'scripts/model_preflight.py', '--config', RUNTIME_CONFIG])
else:
    print('Smoke mode: deterministic fake encoders; outputs are non-scientific.')
'''

SMOKE = r'''# === Cell 3: mandatory CPU smoke test (traverses the whole pipeline in <2 min) ===
if RUN_MODE == 'smoke':
    run_stage('smoke', [sys.executable, 'scripts/run_smoke.py', '--config', 'configs/responsevec.yaml',
              '--workdir', str(ARTIFACT_ROOT / 'smoke')])
    summary = json.loads((ARTIFACT_ROOT / 'smoke' / 'smoke_summary.json').read_text())
    assert summary['status'] == 'PASS', summary
    display(summary)
else:
    print('A100 mode selected; smoke should already have passed in a prior run.')
'''

PREPARE = r'''# === Cell 4: prepare real SocioBench + freeze item folds ===
if RUN_MODE == 'a100':
    repo = PROJECT_ROOT / 'external' / 'SocioBench'
    if not repo.exists():
        repo.parent.mkdir(parents=True, exist_ok=True)
        run('git', 'clone', '--depth', '1', 'https://github.com/JiaWANG-TJ/SocioBench.git', repo)
    run_stage('prepare_data', [sys.executable, 'scripts/prepare_data.py', '--config', RUNTIME_CONFIG])
else:
    print('Skipped real data preparation in smoke mode.')
'''

AUDIT = r'''# === Cell 5: E0/G0 signal audit — THE ZERO-GPU GO/NO-GO ===
# Do not spend A100 hours if sparse histories carry no measurable individual
# signal, or the unseen-item axis has no headroom over the scale-position
# fallback. Stop here and read the report WITH Claude before Cell 6.
if RUN_MODE == 'a100':
    run_stage('signal_audit', [sys.executable, 'scripts/signal_audit.py', '--config', RUNTIME_CONFIG, '--fold', '0'])
    signal = json.loads((ARTIFACT_ROOT / 'metrics' / 'signal_audit.json').read_text())
    display(signal)
    if not (signal['seen_gate_pass'] and signal['unseen_gate_pass']):
        raise RuntimeError('G0 failed. Inspect the audit with Claude before launching 8B extraction.')
'''

ENCODE = r'''# === Cell 6: cache the shared frozen option encoder (once, reused everywhere) ===
if RUN_MODE == 'a100':
    run_stage('encode_options', [sys.executable, 'scripts/encode_options.py', '--config', RUNTIME_CONFIG])
else:
    print('Smoke already validated the option-table contract.')
'''

EXTRACT = r'''# === Cell 7: primary K=5 representation extraction (longest GPU stage) ===
# Cached once for all candidate items, not per fold. Train/validation need only
# canonical option order; final test caches three semantic permutations. Each
# subprocess frees the prior 8B model before loading the next family.
def extract_primary_k(k):
    jobs = [('train', [0]), ('validation', [0]), ('test', OPTION_SEEDS)]
    for split, option_seeds in jobs:
        for option_seed in option_seeds:
            for family in ('causal', 'input_centric', 'response_centric'):
                run_stage(f'extract_{family}_{split}_k{k}_o{option_seed}',
                    [sys.executable, 'scripts/extract_representations.py', '--config', RUNTIME_CONFIG,
                     '--family', family, '--split', split, '--k', str(k),
                     '--option-seed', str(option_seed), '--selection', 'semantic'])

if RUN_MODE == 'a100':
    extract_primary_k(PRIMARY_K)
'''

TRAIN = r'''# === Cell 8: train six outer-fold heads (small; validation selects variants) ===
if RUN_MODE == 'a100':
    option_seed_arg = ','.join(map(str, OPTION_SEEDS))
    decoder_seed_arg = ','.join(map(str, DECODER_SEEDS))
    for fold in OUTER_FOLDS:
        run_stage(f'train_fold{fold}_k{PRIMARY_K}',
            [sys.executable, 'scripts/train_primary.py', '--config', RUNTIME_CONFIG,
             '--fold', str(fold), '--k', str(PRIMARY_K),
             '--option-seeds', option_seed_arg, '--seeds', decoder_seed_arg])
'''

EVALUATE = r'''# === Cell 9: G1 primary decision (crossed bootstrap, Holm family of three) ===
if RUN_MODE == 'a100':
    run_stage('evaluate_primary', [sys.executable, 'scripts/evaluate_primary.py', '--config', RUNTIME_CONFIG, '--k', str(PRIMARY_K)])
    import pandas as pd
    gate = json.loads((ARTIFACT_ROOT / 'metrics' / f'k_{PRIMARY_K}' / 'gate_g1.json').read_text())
    primary = pd.read_csv(ARTIFACT_ROOT / 'metrics' / f'k_{PRIMARY_K}' / 'primary_family_nll.csv')
    display(primary); display(gate)
'''

KCURVE = r'''# === Cell 10: optional K curve (only after inspecting G1) ===
if RUN_MODE == 'a100' and RUN_K_CURVE:
    for k in K_VALUES:
        if k != PRIMARY_K:
            extract_primary_k(k)
        for fold in OUTER_FOLDS:
            run_stage(f'train_fold{fold}_k{k}',
                [sys.executable, 'scripts/train_primary.py', '--config', RUNTIME_CONFIG,
                 '--fold', str(fold), '--k', str(k),
                 '--option-seeds', ','.join(map(str, OPTION_SEEDS)),
                 '--seeds', ','.join(map(str, DECODER_SEEDS))])
        run_stage(f'evaluate_k{k}', [sys.executable, 'scripts/evaluate_primary.py', '--config', RUNTIME_CONFIG, '--k', str(k)])
    run_stage('make_figures', [sys.executable, 'scripts/make_figures.py', '--config', RUNTIME_CONFIG])
'''

ARCHIVE = r'''# === Cell 11: completion record ===
# Negative gate outcomes are results, not files to delete.
if RUN_MODE == 'a100':
    required = [
        ARTIFACT_ROOT / 'processed' / 'item_folds.json',
        ARTIFACT_ROOT / 'metrics' / f'k_{PRIMARY_K}' / 'primary_family_nll.csv',
        ARTIFACT_ROOT / 'metrics' / f'k_{PRIMARY_K}' / 'gate_g1.json',
    ]
    missing = [str(p) for p in required if not p.exists()]
    if missing:
        raise FileNotFoundError(missing)
    make_report('archive', ok=True)
    print('Primary experiment artifacts complete:', ARTIFACT_ROOT)
else:
    print('Smoke complete. Change RUN_MODE to "a100" only after reviewing the smoke summary with Claude.')
'''
def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--notebook", default="output/jupyter-notebook/responsevec_aaai_pipeline.ipynb")
    args = parser.parse_args()
    project = Path(__file__).resolve().parents[1]
    path = (project / args.notebook).resolve()
    path.parent.mkdir(parents=True, exist_ok=True)
    notebook = json.loads(path.read_text(encoding="utf-8")) if path.exists() else {"cells": [], "metadata": {}, "nbformat": 4, "nbformat_minor": 5}

    notebook["cells"] = [
        markdown(INTRO),
        code(MODE),
        markdown("## 1. Setup — extract newest zip, apply Claude patches, resume from Drive backup"),
        code(SETUP),
        code(HELPERS),
        markdown("## 2. Runtime config + (a100) model preflight"),
        code(RUNTIME),
        markdown("## 3. Mandatory CPU smoke test"),
        code(SMOKE),
        markdown("## 4. Prepare real SocioBench and freeze item folds"),
        code(PREPARE),
        markdown("## 5. E0/G0 signal audit — stop and read the report with Claude before spending GPU"),
        code(AUDIT),
        markdown("## 6. Cache the shared option encoder"),
        code(ENCODE),
        markdown("## 7. Primary K=5 representation extraction"),
        code(EXTRACT),
        markdown("## 8. Train six outer-fold heads"),
        code(TRAIN),
        markdown("## 9. G1 primary decision"),
        code(EVALUATE),
        markdown("## 10. Optional K curve, only after inspecting G1"),
        code(KCURVE),
        markdown("## 11. Completion record"),
        code(ARCHIVE),
    ]
    for index, cell in enumerate(notebook["cells"]):
        cell["id"] = f"rv-{index:02d}"
    notebook.setdefault("metadata", {})["kernelspec"] = {
        "display_name": "Python 3", "language": "python", "name": "python3"
    }
    notebook["metadata"]["responsevec"] = {"default_mode": "smoke", "primary_k": 5, "outer_item_folds": 6}
    notebook.setdefault("nbformat", 4)
    notebook.setdefault("nbformat_minor", 5)
    path.write_text(json.dumps(notebook, indent=1), encoding="utf-8")
    print(path)


if __name__ == "__main__":
    main()




